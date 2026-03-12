import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime


def export_documents(answers, run, export_folder):
    doc = Document()

    # Title
    title = doc.add_heading("Healthcare Questionnaire — Completed Answers", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata
    meta = doc.add_paragraph()
    meta.add_run(f"Questionnaire: ").bold = True
    meta.add_run(run.questionnaire_name or "N/A")
    meta.add_run(f"\nGenerated on: ").bold = True
    meta.add_run(datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
    meta.add_run(f"\nTotal Questions: ").bold = True
    meta.add_run(str(len(answers)))
    doc.add_paragraph()

    if not answers:
        doc.add_paragraph("No answers available.")
    else:
        for idx, ans in enumerate(answers, 1):
            # Question
            q_heading = doc.add_heading(f"Q{idx}. {ans.question}", level=2)

            # Answer
            ans_para = doc.add_paragraph()
            ans_para.add_run("Answer: ").bold = True
            ans_para.add_run(ans.answer or "N/A")

            # Confidence
            try:
                conf_val = float(ans.confidence or 0)
                conf_pct = f"{conf_val * 100:.0f}%"
            except Exception:
                conf_pct = "N/A"

            conf_para = doc.add_paragraph()
            conf_para.add_run("Confidence: ").bold = True
            conf_para.add_run(conf_pct)

            # Citation
            cit_para = doc.add_paragraph()
            cit_para.add_run("Source(s): ").bold = True
            cit_para.add_run(ans.citation or "N/A")

            # Evidence snippet
            if ans.evidence and ans.evidence != "N/A":
                ev_para = doc.add_paragraph()
                ev_para.add_run("Evidence Snippet: ").bold = True
                evidence_text = ans.evidence[:300] + "..." if len(ans.evidence) > 300 else ans.evidence
                ev_run = ev_para.add_run(evidence_text)
                ev_run.italic = True
                ev_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

            doc.add_paragraph()

    filename = f"questionnaire_answers_run{run.id}.docx"
    output_path = os.path.join(export_folder, filename)
    doc.save(output_path)
    return output_path
